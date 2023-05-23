from docx import Document
from docx.shared import Inches
import pyttsx3

# Initialize document object
document = Document()

# Speak function to reuse speech method easily
def speak(text):
    '''This function take text as the input and uses the pyttsx3 function to convert the text to speech/sound'''
    pyttsx3.speak(text)

# Profile picture
document.add_picture(
    "Professional_DP.jpg", 
    width=Inches(2.0))

# name, phone number, and email details
speak("What is your name?")
name = input("What is your name? :")
speak(f'{name},What is your phone number?')
phone_number = input("What is your phone number? :")
speak(f'{name},What is your email?')
email = input("What is your email? :")

document.add_paragraph(name + ' | '+ phone_number+' | '+ email)

# about me
data = []
document.add_heading("About me")
speak(f'{name},Tell me about yourself?')
print("Tell me about yourself: ")
while True:
    line = input()
    if line:
        data.append(line)
    else:
        break
finalText = '\n'.join(data)
document.add_paragraph(finalText)

# Work Experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company =  input("Enter company name: ")
from_date = input('From Date: ')
to_date = input('To Date: ')

p.add_run(company +'   ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic=True

experience_details = input('Describe your experience at '+ company + ': ')
p.add_run(experience_details + '\n').italic=True
# More experience details
while True:
    add_more_details = input('Do you want to add more experience details at '+ company + '? Yes or No: ')
    if add_more_details.lower() == 'yes':
        experience_details = input('Include more experience at '+ company + ': ')
        p.add_run(experience_details + '\n').italic=True
    else:
        break

# More experience
while True:
    has_more_experience = input(
        'Do you have more experience? Yes or No: ')
    if has_more_experience.lower() == 'yes':
        # Work Experience
        p = document.add_paragraph()
        # Take experience input from user
        company =  input("Enter company name: ")
        from_date = input('From Date: ')
        to_date = input('To Date: ')

        p.add_run(company +'   ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic=True

        experience_details = input('Describe your experience at '+ company + ': ')
        p.add_run(experience_details).italic=True
        # More experience details
        while True:
            add_more_details = input('Do you want to add more experience details at '+ company + '? Yes or No: ')
            if add_more_details.lower() == 'yes':
                experience_details = input('Include more experience at '+ company + ': ')
                p.add_run(experience_details).italic=True
            else:
                break
    else:
        break

# Add list of Skills
document.add_heading('Skills')
skill = input('Enter skill: ')
# Add skill items
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input(
        'Do you have more skills? Yes or No: ')
    if has_more_skills.lower() == 'yes':
        skill = input('Eneter skill: ')
        # Add skill items
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break
# Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using python program project"

# Save document
document.save('cv.docx')